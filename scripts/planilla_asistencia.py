import calendar
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import os

# Diccionario para traducir meses de español a inglés
meses_es_en = {
    "Enero": "January", "Febrero": "February", "Marzo": "March", "Abril": "April",
    "Mayo": "May", "Junio": "June", "Julio": "July", "Agosto": "August",
    "Septiembre": "September", "Octubre": "October", "Noviembre": "November", "Diciembre": "December"
}

def generar_planilla(dni, nombre, mes, anio=2025):
    if mes not in meses_es_en:
        print("Error: Nombre del mes no válido. Inténtelo de nuevo.")
        return
    
    mes_ingles = meses_es_en[mes]
    num_mes = datetime.strptime(mes_ingles, "%B").month  

    # Crear documento Word
    doc = Document()
    
    # Estilo del título (Arial 14, centrado)
    titulo = doc.add_paragraph()
    run_titulo = titulo.add_run("PLANILLA DE ASISTENCIA PERSONAL MONOTRIBUTISTA\n")
    run_titulo.bold = True
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(14)

    run_subtitulo = titulo.add_run(f"MES DE {mes.upper()} AÑO {anio}")
    run_subtitulo.bold = True
    run_subtitulo.font.name = "Arial"
    run_subtitulo.font.size = Pt(14)
    
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")  # Espaciado
    
    # Datos personales separados en líneas distintas (Arial 12)
    parrafo_nombre = doc.add_paragraph()
    run_nombre = parrafo_nombre.add_run("NOMBRE Y APELLIDO: ")
    run_nombre.bold = True
    run_nombre.font.name = "Arial"
    run_nombre.font.size = Pt(12)
    
    run_nombre_valor = parrafo_nombre.add_run(nombre)
    run_nombre_valor.font.name = "Arial"
    run_nombre_valor.font.size = Pt(12)
        
    parrafo_dni = doc.add_paragraph()
    run_dni = parrafo_dni.add_run("DNI Nº: ")
    run_dni.bold = True
    run_dni.font.name = "Arial"
    run_dni.font.size = Pt(12)
    
    run_dni_valor = parrafo_dni.add_run(dni)
    run_dni_valor.font.name = "Arial"
    run_dni_valor.font.size = Pt(12)
    
    doc.add_paragraph("\n")  # Espaciado
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = 'Table Grid'  # Aplicar bordes a la tabla
    
    # Encabezados de la tabla (Calibri 11, centrado, negrita)
    encabezados = ["Fecha", "Hora de Entrada", "Firma", "Hora de Salida", "Firma", "Total de Horas"]
    hdr_cells = tabla.rows[0].cells
    for i, texto in enumerate(encabezados):
        cell = hdr_cells[i]
        cell.text = texto
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # Agregar filas con datos (Calibri 11)
    dia = 1
    while dia <= calendar.monthrange(anio, num_mes)[1]:
        fecha_actual = datetime(anio, num_mes, dia)
        weekday = fecha_actual.weekday()
        
        row_cells = tabla.add_row().cells
        
        # Si es sábado (5), agregar una sola fila para el fin de semana
        if weekday == 5:
            # Llenar con guiones los campos fecha y total horas
            row_cells[0].text = "-----"  # Fecha
            row_cells[1].text = ""       # Hora de entrada
            row_cells[2].text = "-----"  # Firma entrada
            row_cells[3].text = ""       # Hora de salida
            row_cells[4].text = "-----"  # Firma salida
            row_cells[5].text = "-----"  # Total horas
            dia += 2  # Saltar al lunes
        else:
            # Días laborables (lunes a viernes)
            row_cells[0].text = fecha_actual.strftime("%d/%m/%Y")  # Fecha
            row_cells[1].text = "08:00"  # Hora de entrada
            row_cells[2].text = ""       # Firma entrada
            row_cells[3].text = "14:00"  # Hora de salida
            row_cells[4].text = ""       # Firma salida
            row_cells[5].text = "06:00"  # Total horas
            dia += 1
        
        # Aplicar estilo Calibri 11 a todas las celdas de la fila
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")  # Espaciado
    
    # "Firma Jefe de Unidad" alineado a la derecha (Arial 12)
    firma = doc.add_paragraph()
    firma.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_firma = firma.add_run("Firma Jefe de Unidad")
    run_firma.bold = True
    run_firma.font.name = "Arial"
    run_firma.font.size = Pt(12)

    # Guardar archivo
    nombre_archivo = f"Planilla_{nombre.replace(' ', '_')}_{mes}_{anio}.docx"
    # Crear carpeta si no existe
    carpeta = "planillas-asistencia"
    if not os.path.exists(carpeta):
        os.makedirs(carpeta)

    # Guardar archivo en la carpeta
    nombre_archivo = os.path.join(carpeta, f"Planilla_{nombre.replace(' ', '_')}_{mes}_{anio}.docx")
    doc.save(nombre_archivo)
    print(f"Planilla generada: {nombre_archivo}")